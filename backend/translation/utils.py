import requests
from uuid import UUID
import json
import webvtt
from rest_framework.response import Response
from rest_framework import status
import logging
from docx import *
from docx.shared import Inches
from django.http import HttpResponse, StreamingHttpResponse
from io import StringIO, BytesIO
import os
import datetime
from config import nmt_url, dhruva_key, app_name
from .metadata import LANG_CODE_TO_NAME, english_noise_tags, target_noise_tags
import math
from task.models import Task
import regex
from transcript.utils.timestamp import *
from yt_dlp import YoutubeDL
import pandas as pd
from glossary.tmx.tmxservice import TMXService
from glossary.models import Glossary
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from celery import shared_task
import logging
from django.core.mail import EmailMultiAlternatives
from django.conf import settings

def send_report_as_attachment(subject, body, user, attachment_content, filename, mime_type):
    try:
        msg = EmailMultiAlternatives(
            subject=subject,
            body=body,
            from_email=settings.DEFAULT_FROM_EMAIL,
            to=[user.email],
        )
        
        msg.attach_alternative(body, "text/html")

        msg.attach(filename, attachment_content, mime_type)
        
        msg.send()
        logging.info(f"Document email with attachment '{filename}' sent successfully to {user.email}")
        return True

    except Exception as e:
        logging.error(f"Failed to send email with attachment to {user.email}. Error: {e}")
        return False

def convert_to_scc(subtitles):
    scc_lines = ["Scenarist_SCC V1.0"]

    for index, (timecode, text) in enumerate(subtitles, start=1):
        scc_line = (
            convert_timecode(timecode)
            + "\t94ae 94ae 9420 9420 947a 947a 97a2 97a2 "
            + text_to_hex(text)
            + "92 942c 942c 8080 8080 942f 942f"
        )
        scc_lines.append(scc_line)
    str1 = "\n\n".join(scc_lines)
    return str1


def convert_timecode(timecode):
    parts = timecode.split(" --> ")
    return f"{convert_timestamp(parts[0])}"


def convert_timestamp(timestamp):
    # Convert HH:MM:SS.sss to frames (assuming 30 frames per second)
    hours, minutes, seconds = map(float, timestamp.replace(",", ".").split(":"))
    total_frames = int((hours * 3600 + minutes * 60 + seconds) * 30)
    return f"{total_frames // 1800:02d}:{(total_frames % 1800) // 30:02d}:{(total_frames % 30) * 2:02d}:00"


def text_to_hex(text):
    hex_values = "".join([format(ord(char), "x") for char in text])
    formatted_output = " ".join(
        [hex_values[i : i + 4] for i in range(0, len(hex_values), 4)]
    )
    return formatted_output


def convert_scc_format(payload, task_type):
    if "TRANSCRIPTION" in task_type:
        output_list = [
            (f"{item['start_time']} --> {item['end_time']}", item["text"])
            for item in payload
        ]
    else:
        output_list = [
            (f"{item['start_time']} --> {item['end_time']}", item["target_text"])
            for item in payload
        ]
    scc_content = convert_to_scc(output_list)
    """
    with open("scc_content.txt", "w") as scc_filename:
    	df = pd.DataFrame({'data': [scc_content]})
    	df.to_csv('test.txt', sep='\t', index=False)
    """
    return scc_content


### Utility Functions ###
def validate_uuid4(val):
    try:
        UUID(str(val))
        return True
    except ValueError:
        return False


def valid_xml_char_ordinal(c):
    codepoint = ord(c)
    # conditions ordered by presumed frequency
    return (
        0x20 <= codepoint <= 0xD7FF
        or codepoint in (0x9, 0xA, 0xD)
        or 0xE000 <= codepoint <= 0xFFFD
        or 0x10000 <= codepoint <= 0x10FFFF
    )


def convert_to_docx(content, glossary=""):
    document = Document()
    cleaned_string = "".join(c for c in content if valid_xml_char_ordinal(c))
    paragraph = document.add_paragraph(cleaned_string)

    if glossary != "":
        header = document.add_paragraph("Glossary")
        header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        header_run = header.runs[0]
        header_run.bold = True
        header_run.font.size = Pt(16) 
        table = document.add_table(rows=len(glossary), cols=len(glossary[0]))
        table.style = "Table Grid"

        for row_idx, row in enumerate(glossary):
            for col_idx, value in enumerate(row):
                cell = table.cell(row_idx, col_idx)
                cell.text = value
                paragraph = cell.paragraphs[0]
                run = paragraph.runs[0]
                run.bold = row_idx == 0
                run.font.size = Pt(14 if row_idx == 0 else 12)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 

    run = paragraph.add_run()
    # Prepare document for download
    # -----------------------------
    buffer = BytesIO()
    with open("temp_f.txt", "w") as out_f:
        out_f.write(content)

    buffer.write(open("temp_f.txt", "rb").read())
    print(buffer.seek(0))
    document.save(buffer)
    length = buffer.tell()
    buffer.seek(0)
    response = StreamingHttpResponse(
        streaming_content=buffer,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = "attachment; filename=" + "new_file_download"
    response["Content-Encoding"] = "UTF-8"
    response["Content-Length"] = length
    try:
        os.remove("temp_f.txt")
    except:
        True
    return response

def get_image_from_url(url):
    if not url:
        return None
    try:
        response = requests.get(url, timeout=0.5)
        response.raise_for_status() 
        image_stream = BytesIO(response.content)
        return image_stream
    except requests.exceptions.RequestException as e:
        print(f"Warning: Could not download image from {url}. Error: {e}")
        return None

@shared_task()
def convert_to_paragraph_with_images(payload, video_name, user, task_id, video_d):
    document = Document()
    document.add_paragraph(video_name)
    
    current_text_group = ""
    images_in_group = []
    sentences_in_group = 0
    
    for segment in payload:
        segment_text = ""
        if "verbatim_text" in segment:
            segment_text = segment.get("verbatim_text", "")
        elif "text" in segment:
            segment_text = segment.get("text", "")
        
        if segment_text:
            cleaned_text = segment_text.replace("\n", " ")
            current_text_group += " " + cleaned_text
            sentences_in_group += cleaned_text.count('.')

        if segment.get("image_url"):
            image_stream = get_image_from_url(segment['image_url'])
            if image_stream:
                images_in_group.append(image_stream)
        
        if sentences_in_group >= 5:
            if current_text_group.strip():
                document.add_paragraph(current_text_group.strip())
            
            for img_stream in images_in_group:
                p = document.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = p.add_run()
                run.add_picture(img_stream, width=Inches(5.0))
                img_stream.close()

            document.add_paragraph()
            
            current_text_group = ""
            images_in_group = []
            sentences_in_group = 0

    if current_text_group.strip() or images_in_group:
        if current_text_group.strip():
            document.add_paragraph(current_text_group.strip())
        
        for img_stream in images_in_group:
            p = document.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run()
            run.add_picture(img_stream, width=Inches(5.0))
            img_stream.close()

    buffer = BytesIO()
    document.save(buffer)
    attachment_bytes = buffer.getvalue()
    buffer.close()

    subject = f"Transcription document for task {task_id}"
    email_body = "<p>Your requested report is attached to this email.</p>"
    
    attachment_filename = f"{video_d}.docx"
    docx_mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    send_report_as_attachment(
        subject=subject,
        body=email_body,
        user=user,
        attachment_content=attachment_bytes,
        filename=attachment_filename,
        mime_type=docx_mime_type
    )
    
def convert_to_paragraph(lines, video_name):
    count = 0
    content = ""
    for line in lines:
        content = content + " " + line

    count = 0
    sentences_count = 0
    content = content.replace("\n", " ")
    content = video_name + "\n" + "\n" + content
    for index, i in enumerate(content):
        count += 1
        if content[index] == "." and sentences_count == 5:
            content = content[: index + 1] + "\n" + "\n" + content[index + 1 :]
            sentences_count = 0
        if sentences_count < 5 and i == ".":
            sentences_count += 1

    return content

def convert_to_paragraph_monolingual(payload, video_name, task_id):
    document = Document()
    document.add_paragraph(video_name)

    for segment in payload:
        if segment.get("target_text"):
            document.add_paragraph(segment.get("target_text"))

        if segment.get("image_url"):
            image_stream = get_image_from_url(segment['image_url'])
            if image_stream:
                try:
                    p = document.add_paragraph()
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = p.add_run()
                    run.add_picture(image_stream, width=Inches(5.0))
                    document.add_paragraph()
                except Exception as e:
                    print(f"Warning: Could not add image from {segment['image_url']} to document. Error: {e}")

    buffer = BytesIO()
    document.save(buffer)
    length = buffer.tell()
    buffer.seek(0)

    response = StreamingHttpResponse(
        streaming_content=buffer,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = 'attachment; filename="document_with_images.docx"'
    response["Content-Encoding"] = "UTF-8"
    response["Content-Length"] = length
    
    return response

def convert_to_paragraph_bilingual(payload, video_name, task_id):
    document = Document()
    document.add_paragraph(video_name)
    
    transcripted_content = ""
    translated_content = ""
    images_in_group = []
    sentences_count = 0
    
    for segment in payload:
        if "text" in segment:
            transcripted_content += " " + segment["text"].replace("\n", " ")
        if "target_text" in segment:
            translated_content += " " + segment["target_text"]
        
        if segment.get("image_url"):
            image_stream = get_image_from_url(segment['image_url'])
            if image_stream:
                images_in_group.append(image_stream)

        sentences_count += 1
        
        if sentences_count % 5 == 0:
            if transcripted_content.strip():
                document.add_paragraph(transcripted_content.strip())
            if translated_content.strip():
                document.add_paragraph(translated_content.strip())
            
            for img_stream in images_in_group:
                p = document.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = p.add_run()
                run.add_picture(img_stream, width=Inches(5.0))
                img_stream.close()

            transcripted_content = ""
            translated_content = ""
            images_in_group = []

    if transcripted_content.strip() or translated_content.strip() or images_in_group:
        if transcripted_content.strip():
            document.add_paragraph(transcripted_content.strip())
        if translated_content.strip():
            document.add_paragraph(translated_content.strip())
        
        for img_stream in images_in_group:
            p = document.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run()
            run.add_picture(img_stream, width=Inches(5.0))
            img_stream.close()

    glossary = Glossary.objects.filter(task_ids=task_id)
    if glossary:
        glossary_data = []
        glossary_data.append(["Source Text", "Target Text", "Meaning"])
        for i in glossary:
            glossary_data.append([i.source_text, i.target_text, i.text_meaning or " "])

        document.add_page_break()
        header = document.add_paragraph("Glossary")
        header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        header_run = header.runs[0]
        header_run.bold = True
        header_run.font.size = Pt(16)
        
        table = document.add_table(rows=len(glossary_data), cols=len(glossary_data[0]))
        table.style = "Table Grid"

        for row_idx, row in enumerate(glossary_data):
            for col_idx, value in enumerate(row):
                cell = table.cell(row_idx, col_idx)
                cell.text = str(value)
                paragraph = cell.paragraphs[0]
                run = paragraph.runs[0]
                run.bold = row_idx == 0
                run.font.size = Pt(14 if row_idx == 0 else 12)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 

    buffer = BytesIO()
    document.save(buffer)
    length = buffer.tell()
    buffer.seek(0)
    
    response = StreamingHttpResponse(
        streaming_content=buffer,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = 'attachment; filename="bilingual_document_with_images.docx"'
    response["Content-Encoding"] = "UTF-8"
    response["Content-Length"] = length
    
    return response

def get_batch_translations_using_indictrans_nmt_api(
    sentence_list,
    source_language,
    target_language,
):
    """Function to get the translation for the input sentences using the IndicTrans NMT API.
    Args:
        sentence_list (str): List of sentences to be translated.
        source_language (str): Original language of the sentence.
        target_language (str): Final language of the sentence.
    Returns:
        list: List of dictionaries containing the translated sentences.
    """
    # Convert language code to language text
    source_language_name = LANG_CODE_TO_NAME[source_language]
    target_language_name = LANG_CODE_TO_NAME[target_language]

    # logging.info("source_language_name %s", source_language_name)
    logging.info("target_language_name %s", target_language_name)

    # Create the input sentences list
    input_sentences = [{"source": sentence} for sentence in sentence_list]

    json_data = {
        "input": input_sentences,
        "config": {
            "language": {
                "sourceLanguage": source_language,
                "targetLanguage": target_language,
            },
        },
    }
    try:
        response = requests.post(
            nmt_url,
            headers={"authorization": dhruva_key},
            json=json_data,
        )
        translations_output = response.json()["output"]
        # Collect the translated sentences
        return [translation["target"] for translation in translations_output]
    except Exception as e:
        logging.info("Error in generating translation Output")
        return str(e)


def convert_payload_format(data):
    if data:
        subtitle_url = [item["url"] for item in data if item["ext"] == "vtt"][0]
        subtitle_payload = requests.get(subtitle_url).text
    sentences_list = []
    for vtt_line in webvtt.read_buffer(StringIO(subtitle_payload)):
        start_time = datetime.datetime.strptime(vtt_line.start, "%H:%M:%S.%f")
        unix_start_time = datetime.datetime.timestamp(start_time)
        end_time = datetime.datetime.strptime(vtt_line.end, "%H:%M:%S.%f")
        unix_end_time = datetime.datetime.timestamp(end_time)

        sentences_list.append(
            {
                "start_time": vtt_line.start,
                "end_time": vtt_line.end,
                "text": "",
                "target_text": vtt_line.text,
                "unix_start_time": unix_start_time,
                "unix_end_time": unix_end_time,
            }
        )

    return json.loads(json.dumps({"payload": sentences_list}))


def generate_translation_payload(
    transcript, target_language, list_compare_sources, user_id, url=None
):
    payloads = {}
    if "MACHINE_GENERATED" in list_compare_sources:
        try:
            translation_machine_generated = translation_mg(
                transcript, target_language, user_id
            )
            if type(translation_machine_generated) == Response:
                transcript.transcript_type = "ORIGINAL_SOURCE"
                transcript.save()
                translation_machine_generated = translation_mg(
                    transcript, target_language, user_id
                )
        except:
            if transcript.language == "en":
                transcript.transcript_type = "ORIGINAL_SOURCE"
                transcript.save()
                translation_machine_generated = translation_mg(
                    transcript, target_language, user_id
                )
        transcript.transcript_type = "MACHINE_GENERATED"
        transcript.save()
        payloads["MACHINE_GENERATED"] = translation_machine_generated

    if "MANUALLY_CREATED" in list_compare_sources:
        payload = []
        for txt in transcript.payload["payload"]:
            txt["target_text"] = ""
            payload.append(txt)
        payloads["MANUALLY_CREATED"] = {"payload": payload}
    if "ORIGINAL_SOURCE" in list_compare_sources:
        ydl = YoutubeDL({"format": "best*[acodec!=none]"})
        info = ydl.extract_info(url, download=False)
        subtitles = None
        if "subtitles" in info:
            if target_language in info["subtitles"]:
                subtitles = info["subtitles"][target_language]
                if subtitles is not None:
                    data = convert_payload_format(subtitles)
                    payloads["ORIGINAL_SOURCE"] = data
                else:
                    payloads["ORIGINAL_SOURCE"] = {"payload": []}
            else:
                payloads["ORIGINAL_SOURCE"] = {"payload": []}
        else:
            payloads["ORIGINAL_SOURCE"] = {"payload": []}
    return payloads


def split_at(string, delimiter):
    """
    Desc: Python's split function implementation
    :param string: a string
    :return: a list after breaking string on delimiter match
    """
    result_list = []
    if not delimiter:
        raise ValueError("Empty Separator")

    if not string:
        return [string]
    start = 0
    for index, char in enumerate(string):
        if char == delimiter:
            if not (index < len(string) - 1 and string[index + 1].isdigit()):
                result_list.append(string[start:index])
                start = index + 1
    if start == 0:
        return [string]
    result_list.append(string[start : index + 1])

    return result_list


def get_ratio_of_words(a):
    output = []

    percentage_per_sentence = {}
    pre_last_part = 0
    i = 0
    last = False
    total = 0
    while i < len(a):
        x = split_at(a[i]["text"], ".")
        # list = [len(x)-1]
        # if last == False:
        #     percentage_per_sentence = {}
        percentage_per_sentence[i] = []
        # if(i == 0) handle for no dot and having dots.
        if x[-1] != "":
            first_part = len(x[-1].split())

            j = i + 1
            count_words = 0
            while j != len(a) and len(split_at(a[j]["text"], ".")) == 1:
                count_words += len(a[j]["text"].split())
                j += 1
            last_part = 0
            if j != len(a):
                last_part = len(split_at(a[j]["text"], ".")[0].split())
                # last = True
            # print(x[-1])
            pre_total = total
            total = first_part + count_words + last_part

            if len(x) == 1 and last == False:
                # data_tuple = data_tuple + 0 + first_part*100/total
                # list.append(0)
                # print("i1=", i)
                percentage_per_sentence[i].append(first_part * 100 / total)
            else:
                if last == True:
                    # data_tuple = data_tuple + last_part*100/total + first_part*100/total
                    percentage_per_sentence[i].append(pre_last_part * 100 / pre_total)
                    list_1 = []
                    for l in range(len(x) - 2):
                        list_1.append(100)
                    if len(list_1) > 0:
                        percentage_per_sentence[i].extend(list_1)
                    percentage_per_sentence[i].append(first_part * 100 / total)
                    # print("i2=",i)
                    # output.append(percentage_per_sentence)
                    last = False
                else:
                    # data_tuple = data_tuple + 0 + first_part*100/total
                    # list.append(0)
                    list_1 = []
                    for l in range(len(x) - 1):
                        list_1.append(100)
                    if len(list_1) > 0:
                        percentage_per_sentence[i].extend(list_1)
                    percentage_per_sentence[i].append(first_part * 100 / total)
                    # print("i3=",i)
            # data_tuple = data_tuple + last_part*100/total
            # tuple_list.append(list)

            j = i + 1
            count_words = 0
            while j != len(a) and len(split_at(a[j]["text"], ".")) == 1:
                # print("i4=",j)
                percentage_per_sentence[j] = []
                count_words = len(a[j]["text"].split())
                percentage_per_sentence[j].append(count_words * 100 / total)
                j += 1

            if j != len(a):
                pre_last_part = len(split_at(a[j]["text"], ".")[0].split())
                last = True

            i = j
        else:
            if last == False:
                list_1 = []
                for l in range(len(x) - 1):
                    list_1.append(100)
                # print("i5=",i)
                # percentage_per_sentence[i] = []
                if len(list_1) > 0:
                    percentage_per_sentence[i].extend(list_1)
                # output.append(percentage_per_sentence)
            else:
                percentage_per_sentence[i].append(last_part * 100 / total)
                list_1 = []
                for l in range(len(x) - 2):
                    list_1.append(100)
                if len(list_1) > 0:
                    percentage_per_sentence[i].extend(list_1)
                # print("i6=",i)
                # output.append(percentage_per_sentence)
                last = False
            i += 1
    print(percentage_per_sentence)
    return percentage_per_sentence


def translation_mg(transcript, target_language, user_id, batch_size=25):
    sentence_list = []
    delete_indices = []
    vtt_output = transcript.payload
    ratio_per_sentence = []

    # if transcript.language == "en" and transcript.transcript_type != "ORIGINAL_SOURCE":
    #     ratio_per_sentence = get_ratio_of_words(transcript.payload["payload"])
    #     full_transcript = ""
    #     for index, vtt_line in enumerate(vtt_output["payload"]):
    #         if "text" in vtt_line.keys():
    #             text = vtt_line["text"]
    #             full_transcript = full_transcript + text
    #     sentence_list = split_at(full_transcript, ".")

    #     if sentence_list[-1] == "":
    #         sentence_list.pop()
    #     if sentence_list[-1] == "":
    #         sentence_list.pop()

    # else:
    for index, vtt_line in enumerate(vtt_output["payload"]):
        if "text" in vtt_line.keys():
            text = vtt_line["text"]
            if transcript.language == "en":
                for noise_tag in english_noise_tags:
                    text = text.replace(noise_tag, "")
                sentence_list.append(text)
            else:
                sentence_list.append(text)
            if vtt_line["text"] == "." or vtt_line["text"] == "..":
                delete_indices.append(index)
        else:
            delete_indices.append(index)

    delete_indices.reverse()
    for ind in delete_indices:
        vtt_output["payload"].pop(ind)

    all_translated_sentences = []  # List to store all the translated sentences

    # Iterate over the sentences in batch format and send them to the Translation API
    for i in range(0, len(sentence_list), batch_size):
        batch_of_input_sentences = sentence_list[i : i + batch_size]
        # Get the translation using the Indictrans NMT API
        translations_output = get_batch_translations_using_indictrans_nmt_api(
            sentence_list=batch_of_input_sentences,
            source_language=transcript.language,
            target_language=target_language,
        )
        # Check if translations output doesn't return a string error
        if isinstance(translations_output, str):
            return Response(
                {"message": translations_output}, status=status.HTTP_400_BAD_REQUEST
            )
        else:
            # Add the translated sentences to the list
            all_translated_sentences.extend(translations_output)

    # Check if the length of the translated sentences is equal to the length of the input sentences
    if len(all_translated_sentences) != len(sentence_list):
        return Response(
            {"message": "Error while generating translation."},
            status=status.HTTP_400_BAD_REQUEST,
        )

    if type(ratio_per_sentence) == dict and len(ratio_per_sentence.keys()) > 0:
        sum = 0
        count_sentences = 0
        translated_sentences = {}
        current_sentence_ratio_list = []
        for id, list_ratio in ratio_per_sentence.items():
            for ratio in list_ratio:
                sum = sum + ratio
                current_sentence_ratio_list.append((id, ratio))
                if sum > 99:
                    length_translated_sentence = len(
                        all_translated_sentences[count_sentences]
                    )
                    cleaned_text = regex.sub(
                        r"[^\p{L}\s]", "", all_translated_sentences[count_sentences]
                    ).lower()
                    cleaned_text = regex.sub(
                        r"\s+", " ", cleaned_text
                    )  # for removing multiple blank spaces
                    length_translated_sentence = len(cleaned_text.split(" "))
                    previous_word_index = 0
                    for id, r in current_sentence_ratio_list:
                        number_of_words = math.ceil(
                            (r * length_translated_sentence) / 100
                        )
                        if (
                            previous_word_index + number_of_words
                            < length_translated_sentence
                        ):
                            target_text = " ".join(
                                all_translated_sentences[count_sentences].split(" ")[
                                    previous_word_index : previous_word_index
                                    + number_of_words
                                ]
                            )
                        else:
                            target_text = " ".join(
                                all_translated_sentences[count_sentences].split(" ")[
                                    previous_word_index:
                                ]
                            )

                        if id not in translated_sentences:
                            translated_sentences[id] = target_text
                        else:
                            translated_sentences[id] = (
                                translated_sentences[id] + " " + target_text
                            )
                        previous_word_index = previous_word_index + number_of_words
                    current_sentence_ratio_list = []
                    count_sentences += 1
                    sum = 0

        all_translated_sentences = translated_sentences.values()

    # Update the translation payload with the generated translations
    payload = []
    tmxservice = TMXService()
    import platform

    for source, target in zip(vtt_output["payload"], all_translated_sentences):
        # start_time = datetime.datetime.strptime(source["start_time"], "%H:%M:%S.%f")
        # end_time = datetime.datetime.strptime(source["end_time"], "%H:%M:%S.%f")

        # if platform.system() == "Windows":
        #     # Adjust the start_time and end_time to be within the supported range
        #     unix_start_time = (start_time - datetime.datetime(1970, 1, 1)).total_seconds()
        #     unix_end_time = (end_time - datetime.datetime(1970, 1, 1)).total_seconds()
        # else:
        #     # For other platforms, use the standard timestamp method
        #     unix_start_time = datetime.datetime.timestamp(start_time)
        #     unix_end_time = datetime.datetime.timestamp(end_time)

        start_time = datetime.datetime.strptime(source["start_time"], "%H:%M:%S.%f")
        unix_start_time = datetime.datetime.timestamp(start_time)
        end_time = datetime.datetime.strptime(source["end_time"], "%H:%M:%S.%f")
        unix_end_time = datetime.datetime.timestamp(end_time)

        try:
            if transcript.language == "en":
                noise_tags = list(set(source["text"].split()) & english_noise_tags)
                if noise_tags:
                    replace_noise_tag = target_noise_tags[target_language][
                        noise_tags[0].replace("[", "").replace("]", "")
                    ]
                    if replace_noise_tag != "nan":
                        target = "[" + replace_noise_tag + "] " + target

        except:
            logging.info("Error in replacing noise tags.")

        source["start_time"] = format_timestamp(source["start_time"])
        source["end_time"] = format_timestamp(source["end_time"])

        if "speaker_id" in source.keys():
            locale = transcript.language + "|" + target_language
            org_id = None
            user_id = str(user_id)
            tmx_level = "USER"
            tmx_phrases, res_dict = tmxservice.get_tmx_phrases(
                user_id, org_id, locale, source["text"], tmx_level
            )
            tgt, tmx_replacement = tmxservice.replace_nmt_tgt_with_user_tgt(
                tmx_phrases, source["text"], target
            )
            if len(tmx_replacement) > 0:
                for i in range(len(tmx_replacement)):
                    target = target.replace(
                        tmx_replacement[i]["tgt"], tmx_replacement[i]["tmx_tgt"]
                    )
            payload.append(
                {
                    "start_time": source["start_time"],
                    "end_time": source["end_time"],
                    "text": source["text"],
                    "speaker_id": source["speaker_id"],
                    "unix_start_time": unix_start_time,
                    "unix_end_time": unix_end_time,
                    "target_text": target if source["text"].strip() else source["text"],
                    "image_url": source.get("image_url"),
                }
            )
        else:
            locale = transcript.language + "|" + target_language
            user_id = str(user_id)
            org_id = None
            tmx_level = "USER"
            tmx_phrases, res_dict = tmxservice.get_tmx_phrases(
                user_id, org_id, locale, source["text"], tmx_level
            )
            # [{'src_phrase': 'Python', 'tmx_tgt': 'अजगर', 'tgt': 'पायथन', 'type': 'NMT'}]
            tgt, tmx_replacement = tmxservice.replace_nmt_tgt_with_user_tgt(
                tmx_phrases, source["text"], target
            )
            if len(tmx_replacement) > 0:
                for i in range(len(tmx_replacement)):
                    target = target.replace(
                        tmx_replacement[i]["tgt"], tmx_replacement[i]["tmx_tgt"]
                    )
            payload.append(
                {
                    "start_time": source["start_time"],
                    "end_time": source["end_time"],
                    "text": source["text"],
                    "speaker_id": "",
                    "unix_start_time": unix_start_time,
                    "unix_end_time": unix_end_time,
                    "target_text": target if source["text"].strip() else source["text"],
                    "image_url": source.get("image_url"),
                }
            )
    return json.loads(json.dumps({"payload": payload}))


def set_fail_for_translation_task(task):
    translation_task = (
        Task.objects.filter(target_language=task.target_language)
        .filter(task_type="TRANSLATION_EDIT")
        .filter(video=task.video)
        .first()
    )
    if translation_task is not None:
        translation_task.status = "FAILED"
        translation_task.save()


def convert_to_rt(payload, task_type):
    lines = []
    time_format = "%H:%M:%S.%f"
    lines.append(
        '<Window\n  Width    = "640"\n  Height   = "480"\n  WordWrap = "true"\n  Loop     = "true"\n  bgcolor  = "black"\n>\n<Font\n  Color = "white"\n  Face  = "Arial"\n  Size  = "+2"\n>\n<center>\n<b>\n'
    )
    if "TRANSCRIPTION" in task_type:
        for index, segment in enumerate(payload):
            start_time_str = segment["start_time"]
            end_time_str = segment["end_time"]
            lines.append(
                "<Time begin="
                + f"{start_time_str}"
                + " end="
                + f"{end_time_str}"
                + " />"
                + "<clear/> "
                + " "
                + segment["text"]
            )
    else:
        for index, segment in enumerate(payload):
            start_time_str = segment["start_time"]
            end_time_str = segment["end_time"]
            lines.append(
                "<Time begin="
                + f"{start_time_str}"
                + " end="
                + f"{end_time_str}"
                + " />"
                + "<clear/> "
                + " "
                + segment["target_text"]
            )
    lines.append("</b>\n</center>")
    content = "\n".join(lines)
    return content
